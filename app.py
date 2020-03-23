from fhir_parser import FHIR
from flask import Flask , jsonify, request, render_template, redirect, url_for
import docx
from datetime import datetime

app = Flask(__name__)
fhir = FHIR()
app.config['SECRET_KEY'] = 'secret'



def returnName(uuid):
    patient_uuid = uuid
    test_patient = fhir.get_patient(patient_uuid)
    full_name = test_patient.name.full_name
    return full_name

def returnAddress(uuid):
    patient_uuid = uuid
    test_patient = fhir.get_patient(patient_uuid)
    full_address = test_patient.addresses[0]
    return str(full_address)

def returnNumber(uuid):
    patient_uuid = uuid
    test_patient = fhir.get_patient(patient_uuid)
    telecom = test_patient.telecoms[0]
    return telecom.system , telecom.number
    

def returnALY(uuid):
    patient_uuid = uuid
    test_patient = fhir.get_patient(patient_uuid)   
    DALY = test_patient.get_extension('disability-adjusted-life-years')

    QALY = test_patient.get_extension('quality-adjusted-life-years')
    return DALY , QALY

def returnObservations(uuid):
    patient_uuid = uuid
    test_patient = fhir.get_patient(patient_uuid)
    observations = fhir.get_patient_observations(test_patient.uuid)
    observation_components = []
    for observation in observations:
        observation_components.extend(observation.components)
    observation_component_types = [observation_component.display for observation_component in observation_components]
    observation_component_types.sort()
    #print(str(observation_component_types))
    return (set(observation_component_types))


def makeDocument(name, address, checkedboxes, telecomNum):
    result = [x.strip() for x in address.split(',')]
    doc = docx.Document()
    for p in result:
        doc.add_paragraph(p)
    

    doc.add_paragraph(telecomNum)
    
    doc.add_paragraph(datetime.today().strftime('%Y-%m-%d'))
    
    doc.add_paragraph(" ")
    doc.add_paragraph("Dear " + name)
    doc.add_paragraph("You have been selected by your GP for an appointment regarding your current medical history. Below, your GP has selected a list of observations which they would like to talk about in the next appointment.")
    doc.add_paragraph("Observations to be discussed:")
    for q in checkedboxes:
        doc.add_paragraph(q)
    
    doc.add_paragraph("Please let us know if you have any questions")
    doc.save('ObservationsLetter.docx')



@app.route('/<patientuuid>' , methods=['GET', 'POST'])
def jsonIT(patientuuid):
    telecomSys , telecomNum = returnNumber(patientuuid)
    DALY, QALY = returnALY(patientuuid)  
    name =  str(returnName(patientuuid))
    address = str(returnAddress(patientuuid))
    observations =(returnObservations(patientuuid))
    
    if request.method== "POST" :
        req = request.form 
        checkedboxes =[key for key in req.keys()]
        makeDocument(name, address, checkedboxes, telecomNum)
        return 'Success, a letter has been created and saved called ObservationsLetter.docx!' 


    return render_template('form2.html',name=name, address=address, observations=observations)




@app.route('/form', methods=['GET', 'POST'])
def form():
    #form = Form()
    if request.method== "POST" :
        req = request.form 
        uuid = req["uuids"]
        url = str('/' + uuid)
        return redirect(url)
        
    return render_template("form.html", form=form)